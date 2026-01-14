import streamlit as st
import openai
from docx import Document
import io

# --- Configuración de la Página ---
st.set_page_config(page_title="VoiceToDoc", page_icon="🎙️", layout="centered")

# --- Estilos CSS ---
st.markdown("""
    <style>
    .stButton>button {width: 100%; border-radius: 20px; height: 3em; font-weight: bold;}
    </style>
""", unsafe_allow_html=True)

# --- Funciones ---
def transcribir_y_procesar(api_key, audio_file):
    client = openai.OpenAI(api_key=api_key)
    
    # Transcribir
    if not hasattr(audio_file, 'name') or audio_file.name == 'audio_input':
         audio_file.name = "grabacion.wav"
    transcript = client.audio.transcriptions.create(
        model="whisper-1", file=audio_file, response_format="text"
    )
    
    # Procesar con GPT
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "Eres un redactor experto. Corrige, limpia y da formato profesional al texto."},
            {"role": "user", "content": f"Procesa esto:\n\n{transcript}"}
        ]
    )
    return response.choices[0].message.content

def crear_docx(texto):
    doc = Document()
    doc.add_heading('Documento Transcrito', 0)
    for p in texto.split('\n'):
        if p.strip(): doc.add_paragraph(p)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Interfaz ---
st.title("🎙️ VoiceToDoc")
api_key = st.sidebar.text_input("OpenAI API Key", type="password")

if not api_key:
    st.warning("Ingresa tu API Key en la barra lateral para empezar.")
    st.stop()

tab1, tab2 = st.tabs(["🎤 Grabar", "📂 Subir"])
audio_input = None

with tab1:
    audio_mic = st.audio_input("Grabar")
    if audio_mic: audio_input = audio_mic
with tab2:
    audio_up = st.file_uploader("Subir", type=['mp3', 'wav', 'm4a'])
    if audio_up: audio_input = audio_up

if audio_input and st.button("Transcribir"):
    with st.spinner("Procesando..."):
        try:
            texto = transcribir_y_procesar(api_key, audio_input)
            st.text_area("Resultado", texto, height=300)
            st.download_button("Descargar Word", crear_docx(texto), "transcripcion.docx")
        except Exception as e:
            st.error(f"Error: {e}")
